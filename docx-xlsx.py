import docx
import pandas as pd

# Load the docx file
doc_path = "Title.docx"  # Replace with your DOCX file path
doc = docx.Document(doc_path)

# Create a DataFrame to hold the results
columns = ['Chapter Title', 'Subtitle', 'Subsubtitle', 'REQ']
data = []

# Initialize variables to track the current chapter, subtitle, and subsubtitle
current_chapter = None
current_subtitle = None
current_subsubtitle = None

# Iterate through all paragraphs and tables in the DOCX file
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()

    # Check if the paragraph is a Chapter Title (Heading 1)
    if paragraph.style.name == 'Heading 1':  # Chapter title (Level 1)
        current_chapter = text
        current_subtitle = None
        current_subsubtitle = None  # Reset when a new chapter starts
    # Check if the paragraph is a Subtitle (Heading 2)
    elif paragraph.style.name == 'Heading 2':  # Subtitle (Level 2)
        current_subtitle = text
        current_subsubtitle = None  # Reset subsubtitle when a new subtitle starts
    # Check if the paragraph is a Subsubtitle (Heading 3)
    elif paragraph.style.name == 'Heading 3':  # Subsubtitle (Level 3)
        current_subsubtitle = text

# Iterate through tables and extract REQ from each row
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 1:  # Assuming the first cell contains the REQ
            req = cells[0].text.strip()

            # Append the data to the DataFrame with the current context
            data.append([current_chapter, current_subtitle, current_subsubtitle, req])

# Convert the collected data to a pandas DataFrame
df = pd.DataFrame(data, columns=columns)

# Write the DataFrame to an Excel file
df.to_excel("output_from_docx.xlsx", index=False)

print("Data extraction completed and saved to output_from_docx.xlsx.")
