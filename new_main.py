from docx import Document
import pandas as pd
from pprint import pprint
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from openpyxl.styles import PatternFill
import re
from openpyxl.styles import Border, Side
import win32com.client
import os
# Load the Word document
doc_path = 'Title.docx'
doc = Document(doc_path)
index_merge = []
title_with_tables = []
headings = {}
i = 0
flag_header = False
# Define a global list to store headings with enumeration
global_headings = []

# Helper function to process a table and transform it into a 2-row x 8-column format
def process_table(table):
    rows = table.rows
    if len(rows) < 8 or len(rows[0].cells) < 2:
        return None  # Skip malformed tables
    # Transform the 8x2 table into 2x8
    row_1 = [rows[i].cells[0].text.strip() for i in range(8)]  # First column values
    row_2 = [rows[i].cells[1].text.strip() for i in range(8)]  # Second column values
    return row_1, row_2


# Initialize variables for heading tracking and data storage
heading_1, heading_2, heading_3 = None, None, None
temp_heading = [heading_1, heading_2, heading_3]
requirements_data = []

# Traverse paragraphs and tables in the Word document
for paragraph in doc.paragraphs:
    #print(paragraph._element.xml)  # Print raw XML structure
    if paragraph.style.name == 'Heading 1':
        headings[i] = paragraph.style.name
        index_merge.append(i)
        i = i + 1
        heading_1 = paragraph.text.strip()  # Use the full text if no enumeration


    elif paragraph.style.name == 'Heading 2':
        headings[i] = paragraph.style.name
        index_merge.append(i)
        i = i + 1

        # Extract enumeration and combine with heading text for Heading 2
        match = re.match(r'(\d+(\.\d+)*\.)', paragraph.text.strip())
        if match:
            enumeration = match.group(1)  # Extract the enumeration part
            heading_2 = f"{enumeration} {paragraph.text.strip()}"
        else:
            heading_2 = paragraph.text.strip()  # Use the full text if no enumeration
        #heading_3 = None

    elif paragraph.style.name == 'Heading 3':
        headings[i] = paragraph.style.name
        index_merge.append(i)
        i = i + 1

        # Extract enumeration and combine with heading text for Heading 3
        match = re.match(r'(\d+(\.\d+)*\.)', paragraph.text.strip())
        if match:
            enumeration = match.group(1)  # Extract the enumeration part
            heading_3 = f"{enumeration} {paragraph.text.strip()}"
        else:
            heading_3 = paragraph.text.strip()  # Use the full text if no enumeration

    # write headings
    if heading_1 != temp_heading[0]:
        temp_heading[0] = heading_1
        flag_header = True
        requirements_data.append({
            "col_1": heading_1, "col_2": heading_1,
            "col_3": heading_1, "col_4": heading_1,
            "col_5": heading_1, "col_6": heading_1,
            "col_7": heading_1, "col_8": heading_1})

    if heading_2 != temp_heading[1]:
        temp_heading[1] = heading_2
        flag_header = True
        requirements_data.append({
            "col_1": heading_2, "col_2": heading_2,
            "col_3": heading_2, "col_4": heading_2,
            "col_5": heading_2, "col_6": heading_2,
            "col_7": heading_2, "col_8": heading_2})

    if heading_3 != temp_heading[2]:
        temp_heading[2] = heading_3
        flag_header = True
        requirements_data.append({
            "col_1": heading_3, "col_2": heading_3,
            "col_3": heading_3, "col_4": heading_3,
            "col_5": heading_3, "col_6": heading_3,
            "col_7": heading_3, "col_8": heading_3})

    # Process tables associated with the current heading hierarchy
    if paragraph._element.getnext() is not None and paragraph._element.getnext().tag.endswith('tbl'):
        table_element = paragraph._element.getnext()
        for table in doc.tables:
            if table._element == table_element:
                table_data = process_table(table)
                if table_data:
                    if flag_header == True:
                        flag_header = False
                        requirements_data.append({
                            "col_1": "Requirement name", "col_2": "Statement",
                            "col_3": "Rationale", "col_4": "Add . Info",
                            "col_5": "Maturity", "col_6": "PI Priority",
                            "col_7": "Version", "col_8": " Stakeholder"})
                        i = i + 1
                    row_1, row_2 = table_data
                    requirements_data.append({
                        "col_1": row_2[0], "col_2": row_2[1],
                        "col_3": row_2[2], "col_4": row_2[3],
                        "col_5": row_2[4], "col_6": row_2[5],
                        "col_7": row_2[6], "col_8": row_2[7],
                    })
                    i = i + 1

def extract_headings_with_enumeration(document_path):
    # Access the global variable
    global global_headings

    # Check if the file exists
    if not os.path.exists(document_path):
        raise FileNotFoundError(f"File not found at: {document_path}")

    # Start the Word application
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Keep Word hidden

    # Open the document
    docs = word.Documents.Open(document_path)

    # Extract headings and enumeration
    for paragraph in docs.Paragraphs:
        if paragraph.Range.Style.NameLocal.startswith("Heading"):
            # Extract text and numbering
            text = paragraph.Range.Text.strip()
            numbering = paragraph.Range.ListFormat.ListString  # Extract enumeration
            # Save the result as a combined string to the global list
            global_headings.append(f"{numbering} {text}")

    # Ensure Word and the document are closed properly
    docs.Close(False)
    word.Quit()

# Replace with the path to your Word document
document_path = r"C:\Users\NITRO\PycharmProjects\test\Title.docx"

# Call the function to extract headings
extract_headings_with_enumeration(document_path)

# Print the global list
print("Global Headings List:")
for heading in global_headings:
    print(heading)
# Create a DataFrame
requirements_df = pd.DataFrame(requirements_data)

# Save the processed data to Excel
final_output_path = 'Processed_Requirements_Output.xlsx'
requirements_df.to_excel(final_output_path, index=False)

# Print the output path for confirmation
print(f"Data saved to: {final_output_path}")
pprint({"titles index": index_merge})
print("\n")
pprint(headings)
# Load the existing Excel file
final_output_path = 'Processed_Requirements_Output.xlsx'
wb = load_workbook(final_output_path)
ws = wb.active

ws.delete_rows(1)
for (index, heading), title in zip(headings.items(), global_headings):
    ws.merge_cells(start_row=index + 1, start_column=1, end_row=index + 1, end_column=8)
    ws.cell(row=index + 1, column=1).alignment = Alignment(horizontal="center", vertical="center")

    cell = ws.cell(row=index + 1, column=1)
    if heading == 'Heading 1':
        cell.value = title
        cell.fill = PatternFill(start_color="BDD7EE", end_color="BDD7EE", fill_type="solid")
    elif heading == 'Heading 2':
        cell.value = title
        cell.fill = PatternFill(start_color="F8CBAD", end_color="F8CBAD", fill_type="solid")
    elif heading == 'Heading 3':
        cell.value = title
        cell.fill = PatternFill(start_color="C6E0B4", end_color="C6E0B4", fill_type="solid")

# Define border style
thin_border = Border(
    left=Side(style='thin'),
    right=Side(style='thin'),
    top=Side(style='thin'),
    bottom=Side(style='thin')
)

# Apply borders to all cells in the sheet
for row in ws.iter_rows():
    for cell in row:
        cell.border = thin_border

# Save the updated Excel file
updated_output_path = 'Final_Requirements_Output.xlsx'
wb.save(updated_output_path)

# Confirmation
print(f"Final Excel file saved to: {updated_output_path}")